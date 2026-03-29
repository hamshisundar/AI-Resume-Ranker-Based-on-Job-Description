"""Extract plain text from PDF, DOCX, and legacy DOC (best-effort)."""

from __future__ import annotations

import io
import os
import platform
import shutil
import subprocess
import tempfile

import pdfplumber
from docx import Document

# Resume: first pages only (long CVs); JD PDFs may use a higher limit at call site.
DEFAULT_RESUME_PDF_MAX_PAGES = 5
DEFAULT_JD_PDF_MAX_PAGES = 15


def extract_text_from_pdf_bytes(pdf_bytes: bytes, max_pages: int = DEFAULT_RESUME_PDF_MAX_PAGES) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages[:max_pages]:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts).strip()


def extract_text_from_docx_bytes(data: bytes) -> str:
    buf = io.BytesIO(data)
    doc = Document(buf)
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = " | ".join((c.text or "").strip() for c in row.cells)
            if cells.strip():
                parts.append(cells)
    return "\n".join(parts).strip()


def _extract_doc_via_textutil(data: bytes) -> str | None:
    if platform.system() != "Darwin":
        return None
    path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(data)
            path = tmp.name
        r = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", path],
            capture_output=True,
            text=True,
            timeout=90,
        )
        if r.returncode == 0 and (r.stdout or "").strip():
            return r.stdout.strip()
    except (OSError, subprocess.SubprocessError):
        pass
    finally:
        if path and os.path.isfile(path):
            try:
                os.unlink(path)
            except OSError:
                pass
    return None


def _extract_doc_via_antiword(data: bytes) -> str | None:
    exe = shutil.which("antiword")
    if not exe:
        return None
    path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(data)
            path = tmp.name
        r = subprocess.run(
            [exe, path],
            capture_output=True,
            text=True,
            timeout=90,
        )
        if r.returncode == 0 and (r.stdout or "").strip():
            return r.stdout.strip()
    except (OSError, subprocess.SubprocessError):
        pass
    finally:
        if path and os.path.isfile(path):
            try:
                os.unlink(path)
            except OSError:
                pass
    return None


def extract_text_from_doc_bytes(data: bytes) -> str:
    t = _extract_doc_via_textutil(data)
    if t:
        return t
    t = _extract_doc_via_antiword(data)
    if t:
        return t
    raise ValueError(
        "Could not read .doc. On macOS, TextUtil is used automatically; on Linux install "
        "'antiword', or save the file as .docx or .pdf."
    )


def resume_extension(filename: str) -> str:
    fn = (filename or "").lower().strip()
    for ext in (".docx", ".pdf", ".doc"):
        if fn.endswith(ext):
            return ext
    _, ext = os.path.splitext(fn)
    return ext.lower()


def extract_resume_text(filename: str, data: bytes, *, max_pages_pdf: int = DEFAULT_RESUME_PDF_MAX_PAGES) -> str:
    ext = resume_extension(filename)
    if ext == ".pdf":
        return extract_text_from_pdf_bytes(data, max_pages=max_pages_pdf)
    if ext == ".docx":
        return extract_text_from_docx_bytes(data)
    if ext == ".doc":
        return extract_text_from_doc_bytes(data)
    raise ValueError(f"Unsupported resume type {ext!r}. Use PDF, DOC, or DOCX.")


def jd_is_pdf_filename(filename: str) -> bool:
    return (filename or "").lower().strip().endswith(".pdf")
